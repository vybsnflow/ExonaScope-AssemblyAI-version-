import streamlit as st
import fitz
from docx import Document
from io import BytesIO
from openai import OpenAI
import tempfile
import os
import pytesseract
import requests
import time
from pdf2image import convert_from_bytes
from PIL import Image
import subprocess

# ----- USER AUTH -----
CREDENTIALS = {
    "defenderA": "password123",
    "officeAdmin": "secureBeta!",
    "testUser": "test123"
}

if "authenticated" not in st.session_state:
    st.session_state["authenticated"] = False

if not st.session_state["authenticated"]:
    st.title("🔐 ExonaScope Beta Login")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    if st.button("Login"):
        if username in CREDENTIALS and CREDENTIALS[username] == password:
            st.session_state["authenticated"] = True
        else:
            st.error("Invalid credentials")
    st.stop()

# ----- API CLIENTS -----
client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])
ASSEMBLYAI_API_KEY = os.environ.get("ASSEMBLYAI_API_KEY")

# ----- CASE HEADER -----
st.set_page_config(page_title="ExonaScope – AssemblyAI Edition", layout="centered")
st.title("📂 ExonaScope – Multimodal AI Fact Extractor")

if "case_name" not in st.session_state:
    st.session_state["case_name"] = ""
if "case_number" not in st.session_state:
    st.session_state["case_number"] = ""

with st.form("case_form"):
    st.session_state["case_name"] = st.text_input("Case Name", st.session_state["case_name"])
    st.session_state["case_number"] = st.text_input("Case Number", st.session_state["case_number"])
    if st.form_submit_button("Start New Case"):
        st.session_state.clear()
        st.experimental_rerun()

st.markdown(f"**Case:** {st.session_state['case_name']}  
**Number:** {st.session_state['case_number']}")

# ----- FILE HANDLING -----
uploaded_files = st.file_uploader("Upload PDFs, DOCX, audio, or video", type=["pdf", "docx", "mp3", "wav", "mp4", "m4a"], accept_multiple_files=True)
parsed_segments = []

# ----- PARSING FUNCTIONS -----
def parse_pdf_text(file):
    doc = fitz.open(stream=file.read(), filetype="pdf")
    return "\n".join([page.get_text() for page in doc if page.get_text()])

def run_ocr_on_pdf(file):
    images = convert_from_bytes(file.read(), dpi=300)
    text = ""
    for img in images:
        text += pytesseract.image_to_string(img)
    return text

def parse_docx(file):
    return "\n".join([p.text for p in Document(file).paragraphs])

def transcribe_with_assemblyai(file, suffix):
    headers = {"authorization": ASSEMBLYAI_API_KEY}
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(file.read())
        tmp_path = tmp.name

    with open(tmp_path, "rb") as f:
        upload_url = requests.post("https://api.assemblyai.com/v2/upload", headers=headers, files={"file": f}).json()["upload_url"]

    transcript_req = {
        "audio_url": upload_url,
        "delete_after_seconds": 3600
    }
    transcript_id = requests.post("https://api.assemblyai.com/v2/transcript", headers=headers, json=transcript_req).json()["id"]

    while True:
        poll = requests.get(f"https://api.assemblyai.com/v2/transcript/{transcript_id}", headers=headers).json()
        if poll["status"] == "completed":
            return poll["text"]
        elif poll["status"] == "error":
            return "[Error in transcription]"
        time.sleep(3)

def extract_audio_from_video(file):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as tmp:
        tmp.write(file.read())
        video_path = tmp.name
    audio_path = video_path.replace(".mp4", ".wav")
    subprocess.run(["ffmpeg", "-i", video_path, "-vn", "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1", audio_path], check=True)
    with open(audio_path, "rb") as af:
        result = transcribe_with_assemblyai(af, ".wav")
    os.remove(video_path)
    os.remove(audio_path)
    return result

# ----- PROCESS FILES -----
if uploaded_files:
    st.subheader("📄 Parsed Preview")
    for file in uploaded_files:
        ext = os.path.splitext(file.name)[1].lower()
        st.write(f"📎 File: {file.name} ({file.type})")
        try:
            parsed = ""
            if "pdf" in file.type:
                parsed = parse_pdf_text(file)
                if not parsed.strip():
                    st.info("No embedded text, running OCR...")
                    file.seek(0)
                    parsed = run_ocr_on_pdf(file)
            elif "word" in file.type:
                parsed = parse_docx(file)
            elif "audio" in file.type:
                parsed = transcribe_with_assemblyai(file, ".wav")
            elif "video" in file.type:
                parsed = extract_audio_from_video(file)

            if parsed.strip():
                parsed_segments.append(f"[{file.name}]\n{parsed}")
                with st.expander(f"📃 Preview: {file.name}"):
                    st.text(parsed[:1500])
            else:
                st.warning(f"⚠️ Nothing extractable from: {file.name}")
        except Exception as e:
            st.error(f"❌ Error processing {file.name}: {e}")

# ----- GPT ANALYSIS -----
if parsed_segments and st.button("🧠 Generate Fact Pattern"):
    full_text = "\n\n".join(parsed_segments)
    with st.spinner("Calling GPT-4o..."):
        prompt = f"""You are a legal assistant. Based only on the factual content below, write a chronological, paragraph-based fact pattern suitable for a suppression motion. Do not invent facts. Do not summarize conclusions.

CASE NAME: {st.session_state["case_name"]}
CASE NUMBER: {st.session_state["case_number"]}

SOURCE MATERIAL:
{full_text}
"""
        try:
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You generate legally neutral fact patterns."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2
            )
            result = response.choices[0].message.content.strip()
            st.subheader("📑 Generated Fact Pattern")
            st.text_area("Fact Pattern", value=result, height=300)

            docx_file = BytesIO()
            doc = Document()
            doc.add_heading("Generated Fact Pattern", level=1)
            doc.add_paragraph(result)
            doc.save(docx_file)
            docx_file.seek(0)

            st.download_button("💾 Download (.docx)", docx_file,
                               file_name="fact_pattern.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        except Exception as e:
            st.error(f"❌ GPT Error: {e}")